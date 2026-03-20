import sys
import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

def set_cols(section, n):
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(n))
    cols.set(qn('w:space'), '360')

def keep(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepNext'))

def P(doc, txt, indent=True, italic=False, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=Pt(6), kn=False):
    p = doc.add_paragraph()
    p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Mm(4)
    p.paragraph_format.space_after = after
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(txt)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.italic = italic
    r.bold = bold
    if kn: keep(p)
    return p

def H(doc, txt, lv=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(txt)
    r.font.name = 'Times New Roman'
    if lv == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r.font.size = Pt(10)
    elif lv == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r.font.size = Pt(10)
        r.italic = True
    elif lv == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r.font.size = Pt(10)
        r.bold = True
    keep(p)
    return p

def B(doc, txt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Mm(8)
    p.paragraph_format.first_line_indent = Mm(-4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("\u2022\t" + txt)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    return p

def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Mm(19)
        s.bottom_margin = Mm(25)
        s.left_margin = Mm(15)
        s.right_margin = Mm(15)

    # ── TITLE ──
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_after = Pt(24)
    rt = pt.add_run("RespiriCare: A Digital Healthcare Platform for Respiratory Disease Management through Telehealth and Real-Time Patient Monitoring")
    rt.font.name = 'Times New Roman'; rt.font.size = Pt(18); rt.bold = True

    # ── AUTHORS ──
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER; t.autofit = True
    c1 = t.cell(0,0).paragraphs[0]; c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = c1.add_run("Archana Sasi\nMaster of Computer Applications\nDept. of Computer Applications\nPSG College of Technology\nCoimbatore, India\n23mx118@psgtech.ac.in")
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(11)
    c2 = t.cell(0,1).paragraphs[0]; c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = c2.add_run("Dr. V. Umarani\nAssistant Professor (Sl. Gr.)\nDept. of Computer Applications\nPSG College of Technology\nCoimbatore, India\nvur.mca@psgtech.ac.in")
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)

    doc.add_paragraph()
    ns = doc.add_section(); ns.start_type = 0
    set_cols(ns, 2)

    # ── ABSTRACT ──
    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pa.paragraph_format.space_after = Pt(12)
    ra1 = pa.add_run("Abstract \u2013 "); ra1.font.name = 'Times New Roman'; ra1.font.size = Pt(9.5); ra1.italic = True; ra1.bold = True
    ra2 = pa.add_run(
        "Respiratory diseases such as asthma, chronic obstructive pulmonary disease (COPD), and other breathing-related "
        "conditions require continuous monitoring, timely consultation, and effective communication between patients and "
        "healthcare providers. However, many patients face difficulties in accessing specialized respiratory care due to "
        "geographical limitations, delayed appointments, and lack of digital health support. To address these challenges, "
        "this project proposes RespiriCare, a digital healthcare platform designed to improve respiratory disease management "
        "through telehealth and real-time patient monitoring. RespiriCare integrates key healthcare services including patient "
        "health monitoring, video consultation with doctors, digital prescription management, and pharmacist-assisted "
        "appointment scheduling. In this system, patients can request a video consultation through the pharmacist, who "
        "coordinates with the doctor to schedule the appointment. Once the doctor confirms availability, the pharmacist shares "
        "the consultation link with both the patient and the doctor, enabling seamless remote communication. The platform also "
        "supports digital prescriptions, allowing doctors to generate prescriptions electronically, which can be securely "
        "shared with patients and other healthcare providers. To enhance accessibility, the application provides multi-language "
        "support in English and Tamil, making the system more user-friendly for a wider population. By combining telemedicine, "
        "digital prescription services, and structured communication between patients, pharmacists, and doctors, RespiriCare "
        "aims to create an efficient, accessible, and patient-centered respiratory healthcare solution."
    )
    ra2.font.name = 'Times New Roman'; ra2.font.size = Pt(9.5); ra2.italic = True; ra2.bold = True

    # ── I. INTRODUCTION ──
    H(doc, "I. INTRODUCTION")
    P(doc, "Respiratory diseases represent a significant global health burden, affecting hundreds of millions of people worldwide. "
        "Conditions such as asthma, chronic obstructive pulmonary disease (COPD), bronchitis, and other breathing-related disorders "
        "demand continuous monitoring, periodic medical consultations, and strict adherence to prescribed treatment regimens. The "
        "World Health Organization (WHO) estimates that over 300 million people suffer from asthma alone, while COPD is projected "
        "to become the third leading cause of death globally. Despite the severity of these conditions, many patients in semi-urban "
        "and rural regions continue to face significant barriers in accessing timely respiratory care.")
    P(doc, "Traditional healthcare delivery for respiratory patients relies heavily on periodic in-person clinic visits. Patients "
        "must physically travel to hospitals, wait for appointments, and often face weeks-long delays before receiving specialist "
        "consultations. This delay is particularly dangerous for respiratory conditions, where sudden exacerbations can escalate "
        "rapidly if not addressed within a critical time window. Furthermore, once a prescription is issued, there is no systematic "
        "mechanism to confirm whether the patient has actually procured the medication, understood the dosage, or adhered to the "
        "prescribed schedule.")
    P(doc, "The emergence of mobile health (mHealth) technologies offers a promising solution to these challenges. Smartphones, "
        "now accessible to a vast majority of the global population, provide an ideal platform for continuous health monitoring "
        "and remote clinical communication. However, most existing health applications operate as single-function tools\u2014a separate "
        "application for medication reminders, another for appointment booking, and yet another for accessing prescriptions. This "
        "fragmented landscape forces patients to navigate multiple disconnected platforms, resulting in poor user engagement and "
        "incomplete health tracking.")
    P(doc, "RespiriCare is designed to address this fragmentation by providing a unified digital healthcare platform that connects "
        "the three primary actors in outpatient respiratory care: the Patient, the Doctor, and the Pharmacist. Unlike conventional "
        "applications that treat these roles in isolation, RespiriCare creates a coordinated workflow where each actor contributes "
        "to and benefits from a shared data ecosystem. The patient logs symptoms and tracks medications; the doctor generates "
        "digital prescriptions and conducts teleconsultations; and the pharmacist serves as the scheduling coordinator and "
        "medication verification authority.")
    P(doc, "This paper presents the design, implementation, and evaluation of the RespiriCare platform. Section II outlines the "
        "scope of the application. Section III provides a detailed technology overview. Section IV analyzes existing and proposed "
        "systems. Section V describes the system design including flowcharts. Section VI details individual modules. Section VII "
        "presents comparative analysis. Section VIII concludes with future directions.")

    # ── II. SCOPE ──
    H(doc, "II. SCOPE OF THE APPLICATION")
    P(doc, "The scope of RespiriCare encompasses the complete outpatient respiratory care lifecycle within a single cross-platform "
        "mobile application. The platform specifically targets patients with chronic respiratory conditions who require ongoing "
        "monitoring and periodic professional consultations. The application provides structured role-based access for three "
        "distinct user types: Patients, Doctors, and Pharmacists (who also serve as system administrators).")
    P(doc, "For patients, the platform offers disease selection and tracking, a comprehensive drug directory, symptom logging "
        "with historical views, medication scheduling with reminders, and the ability to request video consultations through the "
        "pharmacist. Patients can also view their prescriptions, track medication adherence, and communicate with healthcare "
        "providers through an integrated chat system.")
    P(doc, "For doctors, the system provides a dedicated dashboard with tools to write digital prescriptions, review patient "
        "cases, access patient details including symptom histories and medication logs, and conduct video consultations. The "
        "prescription module interfaces with a curated drug database to ensure accuracy in medication names, dosages, and "
        "frequencies.")
    P(doc, "For pharmacists, the platform provides administrative capabilities including drug inventory management, prescription "
        "verification, patient notifications, and the critical role of coordinating video consultations between patients and "
        "doctors. The pharmacist receives consultation requests from patients, forwards them to available doctors, collects "
        "availability responses, and shares the meeting link with both parties.")
    P(doc, "Additionally, the application supports multi-language accessibility in English and Tamil, ensuring usability for "
        "a wider population in Southern India. The platform also generates exportable PDF prescriptions and patient reports for "
        "offline reference.")

    # ── III. TECHNOLOGY OVERVIEW ──
    H(doc, "III. TECHNOLOGY OVERVIEW")
    P(doc, "This section outlines the tools, frameworks, and services used in the development of RespiriCare. The technology "
        "stack was selected to maximize cross-platform compatibility, real-time data synchronization, and developer productivity.")

    H(doc, "FLUTTER", lv=3)
    P(doc, "Flutter is an open-source UI software development kit created by Google for building natively compiled applications "
        "for mobile, web, and desktop from a single codebase. Flutter was chosen as the primary framework for RespiriCare because "
        "of its ability to produce high-performance applications for both Android and iOS platforms simultaneously. The framework "
        "uses a widget-based architecture where every visual element is a widget, enabling highly customizable and responsive user "
        "interfaces. Flutter\u2019s hot-reload feature significantly accelerated the development cycle by allowing instant visualization "
        "of code changes without restarting the application.")

    H(doc, "DART", lv=3)
    P(doc, "Dart is a client-optimized programming language developed by Google, serving as the primary language for Flutter "
        "applications. Dart\u2019s sound null safety feature prevents null reference errors at compile time, which is particularly "
        "important in healthcare applications where missing data fields could lead to incorrect medication displays or scheduling "
        "errors. Dart supports both Ahead-of-Time (AOT) compilation for production builds and Just-in-Time (JIT) compilation "
        "for development, providing optimal performance in both contexts.")

    H(doc, "FIREBASE", lv=3)
    P(doc, "Firebase provides the entire backend infrastructure for RespiriCare, eliminating the need for custom server "
        "development and maintenance. The platform leverages several Firebase services:")
    B(doc, "Firebase Authentication: Manages user registration and login, supporting email/password authentication and Google "
        "Sign-In. Firebase Authentication generates and manages secure session tokens, ensuring persistent and safe user sessions "
        "across application restarts.")
    B(doc, "Cloud Firestore: A NoSQL document-oriented database that provides real-time synchronization across all connected "
        "clients. When a doctor writes a prescription, the patient\u2019s device updates instantly without requiring manual refresh. "
        "Firestore\u2019s security rules enforce role-based access control, ensuring patients can only access their own data while "
        "pharmacists have broader administrative access.")
    B(doc, "Firebase Storage: Used for storing profile images and other binary assets uploaded by users through the image picker "
        "functionality.")

    H(doc, "PROVIDER STATE MANAGEMENT", lv=3)
    P(doc, "Provider is a state management solution recommended by the Flutter team. RespiriCare uses Provider to manage "
        "authentication state across the application. The AuthProvider class maintains the current user\u2019s session information "
        "and role, which determines the navigation routing\u2014patients are directed to the Patient Dashboard, doctors to the Doctor "
        "Dashboard, and pharmacists to the Admin Dashboard.")

    H(doc, "ADDITIONAL PACKAGES", lv=3)
    P(doc, "RespiriCare integrates several additional Flutter packages to extend its functionality:")
    B(doc, "flutter_local_notifications and timezone: Provide local push notifications for medication reminders, scheduled "
        "according to the patient\u2019s timezone to ensure timely alerts.")
    B(doc, "url_launcher: Opens external meeting links (such as Google Meet or Jitsi) shared by the pharmacist, enabling "
        "video consultations through established third-party platforms.")
    B(doc, "pdf and path_provider: Generate downloadable PDF prescriptions and patient reports that can be saved locally "
        "or shared externally via the share_plus package.")
    B(doc, "google_fonts: Provides access to the Google Fonts library, ensuring consistent and professional typography "
        "across the application interface.")
    B(doc, "shared_preferences: Stores lightweight local data such as user preferences and language selection settings.")

    # ── IV. SYSTEM ANALYSIS ──
    H(doc, "IV. SYSTEM ANALYSIS")
    P(doc, "System analysis identifies the operational limitations of existing healthcare platforms and establishes the "
        "rationale for the proposed RespiriCare system.")

    H(doc, "A. Existing System", lv=2)
    P(doc, "Current healthcare management for respiratory patients typically involves manual appointment scheduling through "
        "phone calls or hospital reception desks. Prescriptions are handwritten or printed on paper, creating risks of "
        "misinterpretation and loss. There is no standardized mechanism for patients to log daily symptoms or for doctors to "
        "review longitudinal symptom trends between appointments. Medication adherence is tracked only through patient "
        "self-reporting during clinic visits, which is unreliable and subject to recall bias.")
    P(doc, "Video consultation platforms such as Practo, Zocdoc, and generic telemedicine solutions exist but operate "
        "independently from prescription management systems. A patient using one platform for teleconsultation must switch to "
        "a different system to view their prescription, and yet another application to set medication reminders. This "
        "fragmentation results in high dropout rates and poor patient engagement, particularly among elderly patients and "
        "those with limited technological literacy.")

    H(doc, "B. Proposed System", lv=2)
    P(doc, "RespiriCare consolidates all outpatient respiratory care functions into a single application. The proposed system "
        "eliminates the need for multiple disconnected tools by providing integrated symptom tracking, prescription management, "
        "teleconsultation scheduling, and medication reminders within one unified interface. The pharmacist\u2019s role as a "
        "scheduling coordinator is a distinctive feature\u2014rather than requiring direct doctor-patient scheduling (which creates "
        "bottlenecks in doctor availability), the pharmacist acts as an intermediary who collects patient requests, queries "
        "doctor availability, and distributes meeting links to both parties.")
    P(doc, "This triangulated communication model distributes the administrative workload away from the doctor, allowing "
        "physicians to focus exclusively on clinical interactions. The pharmacist, already familiar with the patient\u2019s "
        "prescription history, is uniquely positioned to prioritize consultation requests based on medication adherence "
        "patterns and symptom severity.")

    # ── V. SYSTEM DESIGN ──
    H(doc, "V. SYSTEM DESIGN")
    P(doc, "The system design defines the structural architecture, user workflows, and database schemas that together form "
        "the foundation of RespiriCare.")

    H(doc, "A. System Flowchart", lv=2)
    P(doc, "Figure 5.1 illustrates the complete application flow from user authentication through role-based dashboard "
        "routing to module-specific interactions.", kn=True)

    # ── FLOWCHART IMAGE ──
    fp = "C:/MAD/Therap_app/flowchart.png"
    if os.path.exists(fp):
        try:
            doc.add_picture(fp, width=Inches(3.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Image error: {e}")

    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_after = Pt(12)
    rc = pc.add_run("Figure 5.1 System Flowchart")
    rc.font.name = 'Times New Roman'; rc.font.size = Pt(10)

    P(doc, "The flowchart demonstrates the role-based routing mechanism implemented within the application. Upon launching "
        "the application, users proceed through the Login/SignUp screen. After successful authentication, the system queries "
        "the user\u2019s role from the Firestore database. Based on the role field, the application navigates the user to one of "
        "three distinct dashboards: Patient Dashboard, Doctor Dashboard, or Admin (Pharmacist) Dashboard.")
    P(doc, "The Patient Dashboard provides access to four primary modules: Select Disease (which leads to the Drug Directory "
        "for browsing respiratory medications), My Medications (displaying the patient\u2019s active prescriptions with schedule "
        "tracking and reminders), Log Symptoms (for daily recording of respiratory metrics), and Chat (for messaging healthcare "
        "providers). The medication pathway includes critical sub-flows: setting reminders, taking medicines, and marking "
        "doses as taken or missed.")
    P(doc, "The Doctor Dashboard enables three core functions: Write Rx (creating digital prescriptions), Video Consult "
        "(joining scheduled consultations), and Drug Directory (browsing the medication database). The prescription creation "
        "flow involves composing the prescription, escalating it to the pharmacist for verification, and making it available "
        "to the patient. The consultation flow connects to the Patient Details and Review Case modules, providing the doctor "
        "with comprehensive patient information during the call.")
    P(doc, "The Admin/Pharmacist Dashboard supports Drug Inventory management (updating drug stocks and adding new medications), "
        "Prescription management (verifying, notifying patients, managing prescriptions, and approving requests), and Verification "
        "functions (viewing patient charts, accessing response maps, and reviewing missed medication logs). All three pathways "
        "converge at the END node, representing session completion or logout.")

    H(doc, "B. Database Design", lv=2)
    P(doc, "RespiriCare uses Cloud Firestore\u2019s document-collection model for data storage. The primary collections include "
        "\u2018users\u2019 (storing profile data and role assignments), \u2018prescriptions\u2019 (containing structured medication arrays with "
        "drug name, dosage, frequency, and duration), \u2018consultations\u2019 (tracking request status, timestamps, and meeting links), "
        "\u2018symptom_logs\u2019 (recording daily patient inputs with severity scales), and \u2018medication_logs\u2019 (tracking individual dose "
        "compliance events as taken or missed). Each collection uses the user\u2019s unique ID as a reference key, enabling efficient "
        "role-based querying.")

    H(doc, "C. Consultation Scheduling Workflow", lv=2)
    P(doc, "The consultation scheduling process follows a structured multi-step workflow. The patient initiates a request by "
        "selecting a pharmacist, choosing a preferred date and time, and providing optional notes describing their symptoms. "
        "This request is stored in the \u2018consultations\u2019 collection with a \u2018pending\u2019 status. The pharmacist receives the request "
        "on their dashboard and forwards it to available doctors. The doctor reviews the request and responds with their "
        "availability. The pharmacist then creates a meeting link using an external video platform (such as Google Meet or "
        "Jitsi), confirms the consultation, and the system notifies both the patient and the doctor with the finalized link. "
        "During the scheduled time, both parties join the consultation through the shared link, which opens via the device\u2019s "
        "browser using the url_launcher package.")

    # ── VI. MODULES ──
    H(doc, "VI. MODULES")
    P(doc, "RespiriCare is organized into discrete functional modules, each serving a specific aspect of the respiratory "
        "care workflow.")

    H(doc, "A. User Registration and Authentication", lv=2)
    P(doc, "The registration module supports two authentication methods: email/password registration and Google Sign-In via "
        "the google_sign_in package. New users complete a profile setup that includes selecting their role (patient, doctor, or "
        "pharmacist), provided the role assignment is authorized. Firebase Authentication manages session persistence, "
        "automatically restoring user sessions across application restarts. The login screen implements form validation to "
        "ensure data integrity before submission.")

    H(doc, "B. Disease Selection and Drug Directory", lv=2)
    P(doc, "The disease selection module allows patients to identify their respiratory condition from a curated list. Once a "
        "disease is selected, the system displays relevant drugs from the integrated drug directory. The drug directory "
        "provides comprehensive information about respiratory medications including drug names, compositions, usage instructions, "
        "and side effects. Doctors can also browse this directory when writing prescriptions, ensuring that prescribed "
        "medications are accurately referenced from a standardized database.")

    H(doc, "C. Symptom Logging and History", lv=2)
    P(doc, "The symptom logging module enables patients to record daily respiratory metrics including breathlessness severity, "
        "cough frequency, wheezing intensity, and general well-being. Each log entry is timestamped and stored in the "
        "\u2018symptom_logs\u2019 collection. The symptom history screen presents this data chronologically, allowing both patients and "
        "doctors to identify trends and assess treatment effectiveness over time. Doctors access this information through the "
        "Patient Details screen during or before a consultation.")

    H(doc, "D. Digital Prescription Management", lv=2)
    P(doc, "The digital prescription module enables doctors to generate formatted electronic prescriptions. The Create "
        "Prescription screen provides structured input fields for patient selection, diagnosis, medication selection from the "
        "drug directory, dosage specification, frequency, duration, and additional clinical notes including advice and "
        "follow-up instructions. Completed prescriptions are saved to Firestore and immediately become visible on the patient\u2019s "
        "My Prescriptions screen. The system also generates downloadable PDF versions of prescriptions using the pdf package, "
        "formatted to resemble standard hospital prescription sheets with patient details, doctor information, and a structured "
        "medication table.")

    H(doc, "E. Medication Reminders and Adherence Tracking", lv=2)
    P(doc, "The medication reminder module converts prescription data into scheduled local notifications using the "
        "flutter_local_notifications package. Reminders are configured based on the prescribed frequency and timed according "
        "to the patient\u2019s local timezone using the timezone and flutter_timezone packages. When a notification fires, the "
        "patient can mark the dose as \u2018taken\u2019 or \u2018missed\u2019. This binary event is recorded in the \u2018medication_logs\u2019 collection, "
        "creating a comprehensive adherence timeline. The pharmacist\u2019s dashboard aggregates missed medication events, enabling "
        "proactive follow-up through the Missed Medications screen.")

    H(doc, "F. Teleconsultation via Pharmacist Coordination", lv=2)
    P(doc, "Unlike platforms that implement direct peer-to-peer video calling, RespiriCare adopts a pharmacist-coordinated "
        "teleconsultation model. This design decision reflects the practical reality of clinical environments where doctors "
        "cannot manage their own scheduling alongside patient care. The patient submits a consultation request through the "
        "Request Consultation screen, selecting a pharmacist and preferred time slot. The pharmacist evaluates the request, "
        "coordinates with the doctor to confirm availability, creates a meeting on an external platform, and distributes the "
        "meeting link within the application. Both patient and doctor receive the link and can join the consultation with a "
        "single tap, which opens the link via the device\u2019s browser through the url_launcher package.")

    H(doc, "G. Chat and Communication", lv=2)
    P(doc, "The integrated chat module provides asynchronous text communication between patients, doctors, and pharmacists. "
        "Messages are stored in Firestore and displayed in real-time using Firestore\u2019s snapshot listeners. The chat system "
        "supports individual conversations and allows healthcare providers to send important updates, prescription clarifications, "
        "or follow-up instructions directly to patients without requiring a formal consultation.")

    H(doc, "H. Analytics and Reports", lv=2)
    P(doc, "The analytics module provides visual representations of patient health data through charts and statistical "
        "summaries. The Patient Report screen compiles medication adherence percentages, symptom frequency distributions, and "
        "prescription histories into a structured report that can be exported as a PDF document. Doctors and pharmacists "
        "use this module to quickly assess a patient\u2019s overall health trajectory before or during consultations.")

    # ── VII. COMPARATIVE ANALYSIS ──
    H(doc, "VII. COMPARATIVE ANALYSIS")
    P(doc, "This section compares RespiriCare with existing healthcare management approaches across multiple evaluation "
        "criteria to establish the platform\u2019s relative advantages.")

    H(doc, "A. Communication and Coordination Efficiency", lv=2)
    P(doc, "Traditional healthcare systems require patients to independently contact clinics, wait for callbacks, and "
        "manually coordinate schedules. Many teleconsultation platforms allow direct booking but create scheduling conflicts "
        "as doctors manage both in-person and virtual appointments simultaneously. RespiriCare\u2019s pharmacist-mediated model "
        "eliminates these inefficiencies by introducing a dedicated coordinator who manages the scheduling pipeline, reducing "
        "the administrative burden on doctors and ensuring patients receive confirmed appointment slots with shared meeting links.")
    B(doc, "Patient Convenience: Patients submit a single request and receive a confirmed meeting link without needing to "
        "negotiate availability directly with the doctor.")
    B(doc, "Doctor Efficiency: Doctors focus exclusively on providing medical care, responding to availability queries only "
        "when relevant consultation requests arrive.")

    H(doc, "B. Prescription Accuracy and Accessibility", lv=2)
    P(doc, "Handwritten prescriptions remain the norm in many healthcare facilities, leading to well-documented issues with "
        "legibility, medication errors, and loss of physical documents. Digital prescription systems resolve these issues by "
        "generating standardized, machine-readable prescriptions. RespiriCare further enhances this by integrating the "
        "prescription system with the medication tracking module\u2014when a doctor writes a prescription, the patient\u2019s medication "
        "schedule is automatically populated with reminders, creating a seamless transition from prescription to adherence "
        "tracking.")

    H(doc, "C. Medication Adherence Monitoring", lv=2)
    P(doc, "Conventional adherence tracking relies entirely on patient self-reporting during clinic visits, which is subject "
        "to significant recall bias. Standalone medication reminder applications provide notifications but do not communicate "
        "adherence data back to healthcare providers. RespiriCare bridges this gap by recording every medication event (taken "
        "or missed) in a centralized database accessible to both the prescribing doctor and the supervising pharmacist. This "
        "enables proactive intervention\u2014when the pharmacist notices a pattern of missed medications, they can initiate communication "
        "with the patient through chat or escalate the issue by requesting a follow-up consultation.")

    H(doc, "D. Multi-Language Accessibility", lv=2)
    P(doc, "Healthcare applications designed exclusively in English create barriers for patients in linguistically diverse "
        "regions. RespiriCare addresses this by providing dual-language support in English and Tamil, reflecting the linguistic "
        "demographics of its primary deployment region in Tamil Nadu, Southern India. This localization extends beyond simple "
        "translation to include culturally appropriate medical terminology and interface design, significantly improving "
        "usability for patients with limited English proficiency.")

    H(doc, "E. Cost and Infrastructure Requirements", lv=2)
    P(doc, "RespiriCare\u2019s serverless architecture, powered by Firebase, eliminates the need for dedicated servers, database "
        "administrators, or network infrastructure. Healthcare facilities can deploy the platform by simply distributing the "
        "mobile application to their staff and patients. The Firebase free tier supports substantial usage volumes, making "
        "the platform economically viable for small clinics and community healthcare centers.")

    # ── VIII. CONCLUSION ──
    H(doc, "VIII. CONCLUSION")
    P(doc, "RespiriCare presents a comprehensive digital healthcare solution specifically designed for respiratory disease "
        "management. By integrating patient health monitoring, pharmacist-coordinated teleconsultation scheduling, digital "
        "prescription management, and automated medication adherence tracking into a single cross-platform application, the "
        "system addresses the primary challenges faced by respiratory patients in accessing timely and effective care.")
    P(doc, "The pharmacist-mediated consultation model represents a pragmatic approach to teleconsultation that acknowledges "
        "the administrative realities of clinical practice. Rather than requiring doctors to manage their own scheduling, the "
        "system delegates this responsibility to pharmacists who are already integrated into the prescription workflow, creating "
        "a natural extension of their existing clinical role.")
    P(doc, "Built on Flutter and Firebase, the platform achieves cross-platform compatibility, real-time data synchronization, "
        "and serverless scalability without requiring significant infrastructure investment. The inclusion of Tamil language "
        "support enhances accessibility for the target demographic, while the PDF prescription generation ensures compatibility "
        "with existing offline healthcare documentation workflows.")
    P(doc, "Future enhancements to RespiriCare include integration with wearable respiratory monitoring devices (such as "
        "portable spirometers and pulse oximeters) for automated vital logging, implementation of machine learning algorithms "
        "for symptom pattern recognition and early exacerbation prediction, expansion of language support to include additional "
        "regional languages, and development of a web-based administrative portal for hospital-level deployment and analytics.")

    # ── REFERENCES ──
    H(doc, "REFERENCES")
    refs = [
        "[1]\tGoogle, \u201cFlutter \u2013 Build apps for any screen,\u201d flutter.dev, 2024. [Online]. Available: https://flutter.dev. [Accessed: Mar. 10, 2026].",
        "[2]\tGoogle, \u201cDart programming language,\u201d dart.dev, 2024. [Online]. Available: https://dart.dev. [Accessed: Mar. 10, 2026].",
        "[3]\tGoogle, \u201cFirebase \u2013 App development platform,\u201d firebase.google.com, 2024. [Online]. Available: https://firebase.google.com. [Accessed: Mar. 10, 2026].",
        "[4]\tGoogle, \u201cCloud Firestore documentation,\u201d firebase.google.com/docs/firestore, 2024. [Online]. Available: https://firebase.google.com/docs/firestore. [Accessed: Mar. 10, 2026].",
        "[5]\tWorld Health Organization, \u201cAdherence to long-term therapies: evidence for action,\u201d E. Sabat\u00e9, Ed. WHO, Geneva, Switzerland, 2003.",
        "[6]\tL. Osterberg and T. Blaschke, \u201cAdherence to medication,\u201d New England Journal of Medicine, vol. 353, no. 5, pp. 487\u2013497, Aug. 2005. doi: 10.1056/NEJMra050100.",
        "[7]\tS. Hamine, E. Gerth-Guyette, D. Faulx, B. B. Green, and A. S. Ginsburg, \u201cImpact of mHealth chronic disease management on treatment adherence and patient outcomes: A systematic review,\u201d J. Med. Internet Res., vol. 17, no. 2, e52, Feb. 2015. doi: 10.2196/jmir.3951.",
        "[8]\tGlobal Initiative for Chronic Obstructive Lung Disease, \u201cGlobal Strategy for Prevention, Diagnosis and Management of COPD: 2023 Report,\u201d GOLD, 2023. [Online]. Available: https://goldcopd.org.",
        "[9]\tGoogle, \u201cProvider package for Flutter,\u201d pub.dev, 2024. [Online]. Available: https://pub.dev/packages/provider. [Accessed: Mar. 10, 2026].",
        "[10]\tGoogle, \u201cFirebase Authentication documentation,\u201d firebase.google.com/docs/auth, 2024. [Online]. Available: https://firebase.google.com/docs/auth. [Accessed: Mar. 10, 2026].",
    ]
    for ref in refs:
        pr = doc.add_paragraph()
        pr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pr.paragraph_format.left_indent = Mm(6)
        pr.paragraph_format.first_line_indent = Mm(-6)
        pr.paragraph_format.space_after = Pt(4)
        pr.paragraph_format.line_spacing = 1.1
        rr = pr.add_run(ref)
        rr.font.name = 'Times New Roman'; rr.font.size = Pt(8)

    out = "C:/MAD/Therap_app/RespiriCare_Paper_v2.docx"
    doc.save(out)
    print(f"Saved: {out}")

if __name__ == '__main__':
    build()
